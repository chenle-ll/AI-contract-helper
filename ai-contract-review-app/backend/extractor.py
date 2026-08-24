from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from .config import Settings


class ContractExtractionError(ValueError):
    """Raised when a contract file cannot be parsed."""


async def extract_contract_text(
    file: UploadFile | None,
    pasted_text: str,
    settings: Settings,
) -> tuple[str, int]:
    if file is None:
        text = pasted_text.strip()
        if not text:
            raise ContractExtractionError("未识别到有效合同内容，请检查文件或粘贴文本")
        return text[: settings.max_contract_chars], len(text.encode("utf-8"))

    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in {".docx", ".pdf", ".txt"}:
        raise ContractExtractionError("暂不支持该文件格式，请上传 Word、PDF 或文本文件")

    content = await file.read(settings.max_upload_bytes + 1)
    if len(content) > settings.max_upload_bytes:
        raise ContractExtractionError("文件超过 20MB，请压缩后重新上传")

    try:
        if suffix == ".docx":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
        elif suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            text = content.decode("utf-8-sig")
    except Exception as exc:
        raise ContractExtractionError("文件无法解析，请检查文件是否加密或损坏") from exc

    text = text.replace("\x00", "").strip()
    if not text:
        if suffix == ".pdf":
            raise ContractExtractionError("未识别到 PDF 文本，扫描件请先进行 OCR 后重新上传")
        raise ContractExtractionError("未识别到有效合同内容，请检查文件")
    return text[: settings.max_contract_chars], len(content)
