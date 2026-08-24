from html import escape
from io import BytesIO

from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate

RISK_LABELS = {"high": "高风险", "medium": "中风险", "low": "低风险", "tip": "提示"}
ACTION_LABELS = {
    "pending": "待处理",
    "accept": "已接受",
    "reject": "已拒绝",
    "edit": "已编辑",
    "ignore": "已忽略",
}


def _value(value: object, fallback: str = "—") -> str:
    text = str(value or "").strip()
    return text or fallback


def build_word_report(review: dict) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    document.add_heading("AI 合同审查报告", level=0)
    contract = review.get("contract") or {}
    document.add_paragraph(f"合同名称：{_value(contract.get('name'))}")
    document.add_paragraph(f"合同类型：{_value(contract.get('contractType'))}")
    document.add_paragraph(f"审查立场：{_value(contract.get('reviewStance'))}")
    document.add_paragraph(f"整体风险：{RISK_LABELS.get(review.get('overallRiskLevel'), '—')}")
    document.add_paragraph(f"模型版本：{_value(review.get('modelVersion'))}")

    document.add_heading("风险摘要", level=1)
    document.add_paragraph(_value(review.get("summary")))

    document.add_heading("审查意见", level=1)
    issues = review.get("issues") or []
    if not issues:
        document.add_paragraph("未发现需要列示的审查意见。")
    for index, issue in enumerate(issues, start=1):
        document.add_heading(
            f"{index}. {_value(issue.get('issueType'))} · {RISK_LABELS.get(issue.get('riskLevel'), '—')}",
            level=2,
        )
        fields = [
            ("位置", issue.get("location")),
            ("原文", issue.get("originalText")),
            ("问题", issue.get("reason")),
            ("建议", issue.get("editedText") or issue.get("suggestionText")),
            ("依据", issue.get("basis")),
            ("处理", ACTION_LABELS.get(issue.get("userAction"), _value(issue.get("userAction")))),
        ]
        for label, content in fields:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}：").bold = True
            paragraph.add_run(_value(content))

    missing = review.get("missingClauses") or []
    if missing:
        document.add_heading("缺失条款", level=1)
        for index, item in enumerate(missing, start=1):
            document.add_heading(
                f"{index}. {_value(item.get('name'))} · {RISK_LABELS.get(item.get('riskLevel'), '—')}",
                level=2,
            )
            document.add_paragraph(f"风险：{_value(item.get('reason'))}")
            document.add_paragraph(f"建议：{_value(item.get('suggestion'))}")

    next_steps = review.get("nextSteps") or []
    if next_steps:
        document.add_heading("下一步建议", level=1)
        for item in next_steps:
            document.add_paragraph(_value(item), style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_report(review: dict) -> bytes:
    font_name = "STSong-Light"
    try:
        pdfmetrics.getFont(font_name)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="AI 合同审查报告",
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "ChineseBody", parent=styles["BodyText"], fontName=font_name, fontSize=10,
        leading=16, spaceAfter=5,
    )
    title = ParagraphStyle(
        "ChineseTitle", parent=base, fontSize=21, leading=28, alignment=TA_CENTER, spaceAfter=16,
    )
    heading1 = ParagraphStyle(
        "ChineseH1", parent=base, fontSize=15, leading=22, spaceBefore=12, spaceAfter=7,
    )
    heading2 = ParagraphStyle(
        "ChineseH2", parent=base, fontSize=12, leading=18, spaceBefore=9, spaceAfter=5,
    )

    def paragraph(
        text: object,
        style: ParagraphStyle = base,
        *,
        trusted_markup: bool = False,
    ) -> Paragraph:
        safe = _value(text) if trusted_markup else escape(_value(text))
        safe = safe.replace("\n", "<br/>")
        return Paragraph(safe, style)

    story = [paragraph("AI 合同审查报告", title)]
    contract = review.get("contract") or {}
    metadata = [
        f"合同名称：{_value(contract.get('name'))}",
        f"合同类型：{_value(contract.get('contractType'))}",
        f"审查立场：{_value(contract.get('reviewStance'))}",
        f"整体风险：{RISK_LABELS.get(review.get('overallRiskLevel'), '—')}",
        f"模型版本：{_value(review.get('modelVersion'))}",
    ]
    story.extend(paragraph(item) for item in metadata)
    story.extend([paragraph("风险摘要", heading1), paragraph(review.get("summary"))])
    story.append(paragraph("审查意见", heading1))

    issues = review.get("issues") or []
    if not issues:
        story.append(paragraph("未发现需要列示的审查意见。"))
    for index, issue in enumerate(issues, start=1):
        story.append(paragraph(
            f"{index}. {_value(issue.get('issueType'))} · {RISK_LABELS.get(issue.get('riskLevel'), '—')}",
            heading2,
        ))
        fields = [
            ("位置", issue.get("location")),
            ("原文", issue.get("originalText")),
            ("问题", issue.get("reason")),
            ("建议", issue.get("editedText") or issue.get("suggestionText")),
            ("依据", issue.get("basis")),
            ("处理", ACTION_LABELS.get(issue.get("userAction"), _value(issue.get("userAction")))),
        ]
        for label, content in fields:
            story.append(paragraph(
                f"<b>{escape(label)}：</b>{escape(_value(content))}",
                trusted_markup=True,
            ))

    missing = review.get("missingClauses") or []
    if missing:
        story.append(paragraph("缺失条款", heading1))
        for index, item in enumerate(missing, start=1):
            story.append(paragraph(
                f"{index}. {_value(item.get('name'))} · {RISK_LABELS.get(item.get('riskLevel'), '—')}",
                heading2,
            ))
            story.append(paragraph(f"风险：{_value(item.get('reason'))}"))
            story.append(paragraph(f"建议：{_value(item.get('suggestion'))}"))

    next_steps = review.get("nextSteps") or []
    if next_steps:
        story.append(paragraph("下一步建议", heading1))
        story.extend(paragraph(f"• {_value(item)}") for item in next_steps)

    document.build(story)
    return buffer.getvalue()
